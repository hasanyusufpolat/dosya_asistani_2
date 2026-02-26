"""
PROFESYONEL AKILLI DÜZENLEME MODÜLÜ
Yapay zeka tabanlı içerik düzenleme ve iyileştirme
Tüm dosya türleri için profesyonel belge standardında çıktı üretir
Gelişmiş tipografi, stil koruma ve akıllı biçimlendirme
"""

import os
import re
import logging
import datetime
import hashlib
from typing import Dict, List, Tuple, Optional, Any, Union
from dataclasses import dataclass, field
from enum import Enum
from collections import Counter

# Loglama ayarları
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('ai_editor.log'),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

class EditIntensity(Enum):
    """Düzenleme yoğunluğu seviyeleri"""
    LIGHT = "hafif"        # Sadece boşluklar ve satır sonları
    MEDIUM = "orta"         # Paragraf düzenleme, başlık ekleme
    HEAVY = "yoğun"         # Tam yeniden yapılandırma, tablo oluşturma
    PROFESSIONAL = "profesyonel"  # Kurumsal standartlarda düzenleme

class DocumentStyle(Enum):
    """Belge stilleri"""
    CORPORATE = "kurumsal"      # Şirket içi belgeler
    OFFICIAL = "resmi"           # Resmî yazışmalar
    ACADEMIC = "akademik"        # Tez, makale
    TECHNICAL = "teknik"         # Teknik doküman
    LEGAL = "hukuki"             # Sözleşme, dilekçe
    FINANCIAL = "finansal"       # Rapor, tablo

@dataclass
class DocumentMetrics:
    """Belge metrikleri"""
    total_words: int = 0
    unique_words: int = 0
    avg_word_length: float = 0
    total_sentences: int = 0
    avg_sentence_length: float = 0
    total_paragraphs: int = 0
    avg_paragraph_length: float = 0
    reading_time_minutes: float = 0
    complexity_score: int = 0  # 0-100
    language: str = "unknown"
    has_tables: bool = False
    has_images: bool = False
    has_headings: bool = False
    style_consistency: int = 0  # 0-100
    ocr_confidence: float = 0.0  # OCR güven skoru

@dataclass
class EditResult:
    """Düzenleme sonuç veri yapısı"""
    success: bool
    edited_content: Any
    changes_made: List[str]
    intensity: EditIntensity
    processing_time: float
    word_count: int
    paragraph_count: int
    quality_score: int  # 0-100 arası
    metrics: DocumentMetrics = None
    style: DocumentStyle = DocumentStyle.CORPORATE
    suggestions: List[str] = field(default_factory=list)


# ========== OCR TEMİZLEME YARDIMCI FONKSİYONLARI ==========

def clean_ocr_text(text: str) -> Tuple[str, List[str]]:
    """
    OCR metnini temizle ve gereksiz satırları filtrele
    
    Args:
        text: Orijinal OCR metni
    
    Returns:
        (temizlenmiş_metin, silinen_satırlar)
    """
    lines = text.split('\n')
    cleaned_lines = []
    removed_lines = []
    
    # Gereksiz desenler
    unwanted_patterns = [
        r'\d{1,2}:\d{2}(?::\d{2})?(?:\s*(?:AM|PM|am|pm))?',  # Saat: 14:30, 2:30 PM
        r'\d{1,2}[./]\d{1,2}[./]\d{2,4}',  # Tarih: 15.03.2024
        r'Volte?\s*\d{1,2}:\d{2}',  # Volte 14:30
        r'(?:Turkcell|Vodafone|Türk Telekom|Türkcell)\s*\d{1,2}:\d{2}',
        r'Ekran\s*(?:Alıntısı|Görüntüsü|Fotoğrafı)',  # Ekran Alıntısı
        r'Screen\s*(?:Shot|Capture)',  # Screen Shot/Capture
        r'Screenshot',  # Screenshot
        r'Captur(?:e|ed)',  # Capture/Captured
        r'^\s*\d+\s*$',  # Sadece rakam
        r'^\s*[•\-*]\s*$',  # Sadece madde işareti
        r'^\s*[|\\/]\s*$',  # Sadece çizgi
        r'\d{1,2}\s+(?:Ocak|Şubat|Mart|Nisan|Mayıs|Haziran|Temmuz|Ağustos|Eylül|Ekim|Kasım|Aralık)\s+\d{4}',  # Tarih formatı
        r'\d{1,2}\s+(?:January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{4}',
        r'(?:Sayfa|Page)\s+\d+\s*(?:/|of)\s*\d+',  # Sayfa numaraları
        r'^\s*[_\-\=\*]{3,}\s*$',  # Çizgiler
        r'^\s*[▌▐▀▄█▓▒░]\s*$',  # Blok karakterler
    ]
    
    # Çok kısa satırları filtrele (3 karakterden az)
    for line in lines:
        line = line.strip()
        
        # Boş satırı kontrol et
        if not line:
            cleaned_lines.append('')
            continue
        
        # Çok kısa satırları filtrele
        if len(line) < 3:
            removed_lines.append(line)
            continue
        
        # Gereksiz desenleri kontrol et
        is_unwanted = False
        for pattern in unwanted_patterns:
            if re.search(pattern, line, re.IGNORECASE):
                removed_lines.append(line)
                is_unwanted = True
                break
        
        if not is_unwanted:
            cleaned_lines.append(line)
    
    return '\n'.join(cleaned_lines), removed_lines


def calculate_ocr_confidence(text: str) -> float:
    """
    OCR metninin kalitesini hesapla
    
    Args:
        text: OCR metni
    
    Returns:
        Güven skoru (0-100)
    """
    if not text.strip():
        return 0
    
    words = text.split()
    total_chars = len(text)
    total_words = len(words)
    
    if total_words == 0:
        return 0
    
    # Ortalama kelime uzunluğu (normalde 4-8 arası olmalı)
    avg_word_length = total_chars / total_words
    
    # Anormal kısa kelimeler
    very_short_words = sum(1 for w in words if len(w) <= 2)
    short_ratio = very_short_words / total_words
    
    # Anormal uzun kelimeler (OCR hataları genelde uzun kelimeler oluşturur)
    very_long_words = sum(1 for w in words if len(w) > 20)
    long_ratio = very_long_words / total_words
    
    # Özel karakter oranı
    special_chars = sum(1 for c in text if not c.isalnum() and not c.isspace())
    special_ratio = special_chars / total_chars if total_chars > 0 else 0
    
    # Rakam oranı (çok fazla rakam OCR hatası olabilir)
    digit_chars = sum(1 for c in text if c.isdigit())
    digit_ratio = digit_chars / total_chars if total_chars > 0 else 0
    
    # Büyük harf oranı (çok fazla büyük harf OCR hatası olabilir)
    uppercase_chars = sum(1 for c in text if c.isupper())
    uppercase_ratio = uppercase_chars / total_chars if total_chars > 0 else 0
    
    # Güven skoru hesapla
    confidence = 100
    
    # Ortalama kelime uzunluğu çok düşükse
    if avg_word_length < 3:
        confidence -= 30
    elif avg_word_length < 4:
        confidence -= 15
    elif avg_word_length > 15:
        confidence -= 20
    
    # Çok fazla kısa kelime varsa
    if short_ratio > 0.3:
        confidence -= 25
    elif short_ratio > 0.2:
        confidence -= 15
    
    # Çok fazla uzun kelime varsa
    if long_ratio > 0.1:
        confidence -= 20
    elif long_ratio > 0.05:
        confidence -= 10
    
    # Çok fazla özel karakter varsa
    if special_ratio > 0.2:
        confidence -= 20
    elif special_ratio > 0.1:
        confidence -= 10
    
    # Çok fazla rakam varsa (sayısal belgeler hariç)
    if digit_ratio > 0.3:
        confidence -= 15
    elif digit_ratio > 0.2:
        confidence -= 5
    
    # Çok fazla büyük harf varsa
    if uppercase_ratio > 0.5:
        confidence -= 15
    elif uppercase_ratio > 0.3:
        confidence -= 5
    
    return max(0, min(100, confidence))


def merge_intelligent_lines(text: str) -> str:
    """
    Akıllı satır birleştirme - bölünmüş kelimeleri ve cümleleri düzelt
    
    Args:
        text: Ham metin
    
    Returns:
        Birleştirilmiş metin
    """
    lines = text.split('\n')
    merged = []
    i = 0
    
    while i < len(lines):
        current_line = lines[i].strip()
        
        # Boş satır
        if not current_line:
            merged.append('')
            i += 1
            continue
        
        # Sonraki satır var mı?
        if i < len(lines) - 1:
            next_line = lines[i+1].strip()
            
            # Eğer sonraki satır küçük harfle başlıyorsa ve
            # mevcut satır noktalama işareti ile bitmiyorsa
            if next_line and next_line[0].islower() and not current_line[-1] in '.!?:;,':
                # Birleştir
                if current_line.endswith('-'):
                    # Kelime bölünmesi
                    current_line = current_line[:-1] + next_line
                else:
                    current_line += ' ' + next_line
                i += 2  # Bir sonraki satırı atla
                continue
        
        merged.append(current_line)
        i += 1
    
    return '\n'.join(merged)


def normalize_whitespace(text: str) -> str:
    """
    Boşlukları normalize et
    
    Args:
        text: Ham metin
    
    Returns:
        Düzenlenmiş metin
    """
    # Fazla boşlukları temizle
    text = re.sub(r' +', ' ', text)
    
    # Fazla satır sonlarını temizle
    text = re.sub(r'\n\s*\n\s*\n', '\n\n', text)
    
    return text.strip()


def fix_common_ocr_errors(text: str) -> str:
    """
    Yaygın OCR hatalarını düzelt
    
    Args:
        text: OCR metni
    
    Returns:
        Düzeltilmiş metin
    """
    # Yaygın hata düzeltmeleri
    replacements = [
        # Sayısal hatalar
        (r'0', 'O', r'\b0\b'),  # Tek başına 0 -> O
        (r'1', 'l', r'\b1\b'),  # Tek başına 1 -> l
        (r'5', 'S', r'\b5\b'),  # Tek başına 5 -> S
        
        # Karakter karışıklıkları
        (r'§', 'S'),  # Paragraf işareti -> S
        (r'©', 'C'),  # Copyright -> C
        (r'®', 'R'),  # Registered -> R
        (r'™', 'TM'),  # Trademark -> TM
        (r'•', '-'),  # Madde işareti -> -
        (r'·', '.'),  # Orta nokta -> .
        (r'…', '...'),  # Üç nokta -> ...
        (r'–', '-'),  # Uzun tire -> -
        (r'—', '-'),  # Uzun tire -> -
        (r'"', '"'),  # Tırnak işareti
        (r'"', '"'),  # Tırnak işareti
        
        # Türkçe karakter düzeltmeleri
        (r'Ý', 'İ'),  # Y. harfi -> İ
        (r'Þ', 'Ş'),  # TH -> Ş
        (r'ð', 'ğ'),  # Ð -> ğ
        (r'ý', 'ı'),  # ý -> ı
    ]
    
    for old, new in replacements:
        if isinstance(old, tuple):
            # Regex pattern
            text = re.sub(old[0], new, text)
        else:
            # Direkt değiştirme
            text = text.replace(old, new)
    
    return text


class AIEditor:
    """Profesyonel Akıllı Düzenleyici - Gelişmiş Versiyon"""
    
    def __init__(self):
        self.supported_formats = {
            'docx': self.edit_word,
            'doc': self.edit_word,
            'txt': self.edit_text,
            'pdf': self.edit_pdf,
            'png': self.edit_image,
            'jpg': self.edit_image,
            'jpeg': self.edit_image,
            'rtf': self.edit_text,
            'md': self.edit_markdown
        }
        
        # Profesyonel belge standartları (gelişmiş)
        self.professional_standards = {
            'max_paragraph_length': 500,      # karakter
            'min_paragraph_length': 30,        # karakter
            'max_sentence_length': 30,          # kelime
            'min_sentence_length': 3,           # kelime
            'max_line_length': 80,              # karakter (terminal standardı)
            'heading_patterns': [
                r'^#{1,6}\s+',                  # Markdown başlıkları
                r'^[A-Z][A-Z\s]{2,}$',          # TAMAMEN BÜYÜK HARF
                r'^[0-9]+\.\s+',                 # Numaralı başlıklar
                r'^[•\-*]\s+',                    # Madde işaretleri
                r'^[A-Za-zğüşıöçĞÜŞİÖÇ]\.\s+',    # Harfli listeler
            ],
            'table_indicators': ['\t', '|', '  '],
            'list_indicators': ['•', '-', '*', '→', '✓', '✅', '●', '○'],
            'numeric_indicators': [r'^\d+\.', r'^\d+\)', r'^\(\d+\)']
        }
        
        # Dil desenleri
        self.language_patterns = {
            'tr': {
                'chars': set('ğüşıöçĞÜŞİÖÇ'),
                'words': ['ve', 'bir', 'bu', 'için', 'ile', 'olarak', 'ancak', 'veya', 'çünkü'],
                'stopwords': ['acaba', 'ama', 'aslında', 'az', 'bazı', 'belki', 'biri', 'birkaç', 
                             'birşey', 'biz', 'bu', 'çok', 'çünkü', 'da', 'daha', 'de', 'defa', 
                             'diye', 'eğer', 'en', 'gibi', 'hem', 'hep', 'hepsi', 'her', 'hiç', 
                             'için', 'ile', 'ise', 'kez', 'ki', 'kim', 'mı', 'mu', 'mü', 'nasıl', 
                             'ne', 'neden', 'nerde', 'nerede', 'nereye', 'niçin', 'niye', 'o', 
                             'sanki', 'şey', 'siz', 'şu', 'tüm', 've', 'veya', 'ya', 'yani']
            },
            'en': {
                'chars': set(),
                'words': ['the', 'and', 'for', 'with', 'this', 'that', 'from', 'have', 'will'],
                'stopwords': ['a', 'about', 'above', 'after', 'again', 'against', 'all', 'am', 
                             'an', 'and', 'any', 'are', "aren't", 'as', 'at', 'be', 'because', 
                             'been', 'before', 'being', 'below', 'between', 'both', 'but', 'by', 
                             "can't", 'cannot', 'could', "couldn't", 'did', "didn't", 'do', 'does',
                             "doesn't", 'doing', "don't", 'down', 'during', 'each', 'few', 'for', 
                             'from', 'further', 'had', "hadn't", 'has', "hasn't", 'have', "haven't",
                             'having', 'he', "he'd", "he'll", "he's", 'hence', 'her', 'here', 
                             "here's", 'hers', 'herself', 'him', 'himself', 'his', 'how', "how's",
                             'i', "i'd", "i'll", "i'm", "i've", 'if', 'in', 'into', 'is', "isn't",
                             'it', "it's", 'its', 'itself', "let's", 'me', 'more', 'most', "mustn't",
                             'my', 'myself', 'no', 'nor', 'not', 'of', 'off', 'on', 'once', 'only',
                             'or', 'other', 'ought', 'our', 'ours', 'ourselves', 'out', 'over',
                             'own', 'same', "shan't", 'she', "she'd", "she'll", "she's", 'should',
                             "shouldn't", 'so', 'some', 'such', 'than', 'that', "that's", 'the',
                             'their', 'theirs', 'them', 'themselves', 'then', 'there', "there's",
                             'therefore', 'these', 'they', "they'd", "they'll", "they're", "they've",
                             'this', 'those', 'though', 'through', 'to', 'too', 'under', 'until',
                             'up', 'very', 'was', "wasn't", 'we', "we'd", "we'll", "we're", "we've",
                             'were', "weren't", 'what', "what's", 'when', "when's", 'where', "where's",
                             'which', 'while', 'who', "who's", 'whom', 'why', "why's", 'with',
                             "won't", 'would', "wouldn't", 'you', "you'd", "you'll", "you're", "you've",
                             'your', 'yours', 'yourself', 'yourselves']
            }
        }
        
        # Stil kuralları
        self.style_rules = {
            DocumentStyle.CORPORATE: {
                'heading_format': '### {}',
                'list_format': '• {}',
                'emphasis': '**{}**',
                'alignment': 'left',
                'line_spacing': 1.15,
                'paragraph_spacing': 1.5,
                'font_family': 'Calibri',
                'font_size': 11
            },
            DocumentStyle.OFFICIAL: {
                'heading_format': '{}\n' + '-' * 40,
                'list_format': '{}. {}',
                'emphasis': '_{}_',
                'alignment': 'justify',
                'line_spacing': 1.5,
                'paragraph_spacing': 2.0,
                'font_family': 'Times New Roman',
                'font_size': 12
            },
            DocumentStyle.ACADEMIC: {
                'heading_format': '{}\n' + '=' * 40,
                'list_format': '{}. {}',
                'emphasis': '*{}*',
                'alignment': 'justify',
                'line_spacing': 2.0,
                'paragraph_spacing': 2.5,
                'font_family': 'Arial',
                'font_size': 11
            },
            DocumentStyle.TECHNICAL: {
                'heading_format': '## {}',
                'list_format': '- {}',
                'emphasis': '`{}`',
                'alignment': 'left',
                'line_spacing': 1.0,
                'paragraph_spacing': 1.2,
                'font_family': 'Consolas',
                'font_size': 10
            },
            DocumentStyle.LEGAL: {
                'heading_format': '{}',
                'list_format': '{}. {}',
                'emphasis': '_{}_',
                'alignment': 'justify',
                'line_spacing': 1.5,
                'paragraph_spacing': 2.0,
                'font_family': 'Garamond',
                'font_size': 12
            },
            DocumentStyle.FINANCIAL: {
                'heading_format': '{}',
                'list_format': '• {}',
                'emphasis': '**{}**',
                'alignment': 'right' if 'tutar' in str else 'left',
                'line_spacing': 1.0,
                'paragraph_spacing': 1.2,
                'font_family': 'Arial',
                'font_size': 10
            }
        }
    
    # ========== METİN ANALİZ FONKSİYONLARI ==========
    
    def analyze_text_metrics(self, text: str) -> DocumentMetrics:
        """Metin metriklerini analiz et"""
        metrics = DocumentMetrics()
        
        if not text.strip():
            return metrics
        
        # Kelime sayısı
        words = re.findall(r'\b\w+\b', text.lower())
        metrics.total_words = len(words)
        metrics.unique_words = len(set(words))
        
        if metrics.total_words > 0:
            # Ortalama kelime uzunluğu
            metrics.avg_word_length = sum(len(w) for w in words) / metrics.total_words
            
            # Cümle sayısı
            sentences = re.split(r'[.!?]+', text)
            metrics.total_sentences = len([s for s in sentences if s.strip()])
            
            if metrics.total_sentences > 0:
                metrics.avg_sentence_length = metrics.total_words / metrics.total_sentences
            
            # Paragraf sayısı
            paragraphs = text.split('\n\n')
            metrics.total_paragraphs = len([p for p in paragraphs if p.strip()])
            
            if metrics.total_paragraphs > 0:
                metrics.avg_paragraph_length = metrics.total_words / metrics.total_paragraphs
            
            # Okuma süresi (dakika) - ortalama 200 kelime/dakika
            metrics.reading_time_minutes = metrics.total_words / 200
            
            # Karmaşıklık skoru
            complexity = 0
            complexity += min(100, metrics.avg_sentence_length * 3)  # Uzun cümleler
            complexity += min(100, metrics.avg_word_length * 10)     # Uzun kelimeler
            complexity += len([w for w in words if len(w) > 10])      # Çok uzun kelimeler
            metrics.complexity_score = min(100, complexity)
            
            # Dil tespiti
            tr_count = sum(1 for c in text if c in self.language_patterns['tr']['chars'])
            en_count = sum(1 for w in words if w in self.language_patterns['en']['words'])
            
            if tr_count > en_count * 2:
                metrics.language = 'tr'
            elif en_count > tr_count:
                metrics.language = 'en'
            else:
                metrics.language = 'mixed'
        
        return metrics
    
    def detect_style(self, text: str, metrics: DocumentMetrics) -> DocumentStyle:
        """Belge stilini tespit et"""
        text_lower = text.lower()
        
        # Stil tespiti için anahtar kelimeler
        style_keywords = {
            DocumentStyle.CORPORATE: ['şirket', 'firma', 'kurum', 'departman', 'proje', 'rapor', 
                                      'toplantı', 'müdür', 'genel', 'kurumsal'],
            DocumentStyle.OFFICIAL: ['resmi', 'makam', 'bakanlık', 'valilik', 'kaymakamlık',
                                     'belediye', 'idare', 'tebliğ', 'karar', 'yazı'],
            DocumentStyle.ACADEMIC: ['tez', 'makale', 'araştırma', 'çalışma', 'inceleme',
                                     'üniversite', 'fakülte', 'öğrenci', 'akademik', 'bilimsel'],
            DocumentStyle.TECHNICAL: ['teknik', 'sistem', 'yazılım', 'donanım', 'kod',
                                      'veri', 'tablo', 'grafik', 'parametre', 'fonksiyon'],
            DocumentStyle.LEGAL: ['sözleşme', 'madde', 'fıkra', 'kanun', 'yönetmelik',
                                  'hüküm', 'taraf', 'ibraz', 'taahhüt', 'yemin'],
            DocumentStyle.FINANCIAL: ['banka', 'dekont', 'havale', 'eft', 'ödeme', 'tutar',
                                     'bakiye', 'hesap', 'alacak', 'borç', 'cari', 'fatura']
        }
        
        scores = {}
        for style, keywords in style_keywords.items():
            score = sum(1 for k in keywords if k in text_lower)
            if score > 0:
                scores[style] = score
        
        if scores:
            return max(scores, key=scores.get)
        
        # Varsayılan stil
        if metrics.total_words > 1000:
            return DocumentStyle.ACADEMIC
        elif metrics.has_tables:
            return DocumentStyle.FINANCIAL
        else:
            return DocumentStyle.CORPORATE
    
    # ========== GELİŞMİŞ DÜZENLEME FONKSİYONLARI ==========
    
    def clean_text_advanced(self, text: str, metrics: DocumentMetrics) -> Tuple[str, List[str]]:
        """Gelişmiş metin temizleme"""
        changes = []
        original_lines = text.split('\n')
        cleaned_lines = []
        
        i = 0
        while i < len(original_lines):
            line = original_lines[i].rstrip()
            
            # Boş satırları düzenle
            if not line.strip():
                # Birden fazla boş satır varsa tekille
                if i > 0 and cleaned_lines and not cleaned_lines[-1].strip():
                    i += 1
                    continue
                cleaned_lines.append('')
                i += 1
                continue
            
            # Satır sonundaki boşlukları temizle
            line = line.rstrip()
            
            # Çok uzun satırları böl
            if len(line) > self.professional_standards['max_line_length']:
                # Kelimelere böl
                words = line.split()
                current_line = []
                current_length = 0
                
                for word in words:
                    if current_length + len(word) + 1 <= self.professional_standards['max_line_length']:
                        current_line.append(word)
                        current_length += len(word) + 1
                    else:
                        if current_line:
                            cleaned_lines.append(' '.join(current_line))
                            changes.append(f"Uzun satır bölündü")
                        current_line = [word]
                        current_length = len(word)
                
                if current_line:
                    cleaned_lines.append(' '.join(current_line))
            else:
                cleaned_lines.append(line)
            
            i += 1
        
        return '\n'.join(cleaned_lines), changes
    
    def format_paragraphs(self, text: str, metrics: DocumentMetrics) -> Tuple[str, List[str]]:
        """Paragrafları formatla"""
        changes = []
        paragraphs = text.split('\n\n')
        formatted_paragraphs = []
        
        for para in paragraphs:
            if not para.strip():
                continue
            
            # Paragraftaki satırları birleştir
            lines = para.split('\n')
            if len(lines) > 1:
                # Başlık kontrolü
                if any(para.startswith(p) for p in self.professional_standards['heading_patterns']):
                    formatted_paragraphs.append(para)
                else:
                    # Normal paragraf
                    merged = ' '.join(line.strip() for line in lines)
                    formatted_paragraphs.append(merged)
                    if len(lines) > 1:
                        changes.append("Paragraf birleştirildi")
            else:
                formatted_paragraphs.append(para)
        
        return '\n\n'.join(formatted_paragraphs), changes
    
    def detect_and_format_lists(self, text: str) -> Tuple[str, List[str]]:
        """Listeleri tespit et ve formatla"""
        changes = []
        lines = text.split('\n')
        formatted_lines = []
        in_list = False
        list_items = []
        
        for i, line in enumerate(lines):
            stripped = line.strip()
            
            # Liste öğesi kontrolü
            is_list_item = False
            list_type = None
            
            for indicator in self.professional_standards['list_indicators']:
                if stripped.startswith(indicator):
                    is_list_item = True
                    list_type = 'bullet'
                    break
            
            if not is_list_item:
                for pattern in self.professional_standards['numeric_indicators']:
                    if re.match(pattern, stripped):
                        is_list_item = True
                        list_type = 'numeric'
                        break
            
            if is_list_item:
                if not in_list:
                    in_list = True
                    list_items = [stripped]
                else:
                    list_items.append(stripped)
            else:
                if in_list:
                    # Listeyi formatla
                    for idx, item in enumerate(list_items, 1):
                        if list_type == 'bullet':
                            formatted_lines.append(f"• {item.lstrip('•-*→✓✅●○').strip()}")
                        else:
                            # Numaralı listeyi düzelt
                            content = re.sub(r'^\d+[.)]', '', item).strip()
                            formatted_lines.append(f"{idx}. {content}")
                    changes.append(f"{len(list_items)} öğeli liste formatlandı")
                    in_list = False
                    list_items = []
                    formatted_lines.append(line)
                else:
                    formatted_lines.append(line)
        
        # Son listeyi ekle
        if in_list:
            for idx, item in enumerate(list_items, 1):
                if list_type == 'bullet':
                    formatted_lines.append(f"• {item.lstrip('•-*→✓✅●○').strip()}")
                else:
                    content = re.sub(r'^\d+[.)]', '', item).strip()
                    formatted_lines.append(f"{idx}. {content}")
            changes.append(f"{len(list_items)} öğeli liste formatlandı")
        
        return '\n'.join(formatted_lines), changes
    
    def detect_and_format_tables(self, text: str) -> Tuple[str, List[str], bool]:
        """Tabloları tespit et ve formatla"""
        changes = []
        has_tables = False
        lines = text.split('\n')
        formatted_lines = []
        
        i = 0
        while i < len(lines):
            line = lines[i]
            
            # Tablo satırı kontrolü
            if any(indicator in line for indicator in self.professional_standards['table_indicators']):
                # Tablo başlangıcı
                table_rows = []
                while i < len(lines) and any(indicator in lines[i] for indicator in self.professional_standards['table_indicators']):
                    table_rows.append(lines[i])
                    i += 1
                
                if len(table_rows) > 1:
                    # Tabloyu işle
                    formatted_table = self._format_table(table_rows)
                    formatted_lines.append(formatted_table)
                    changes.append(f"{len(table_rows)} satırlı tablo formatlandı")
                    has_tables = True
                else:
                    formatted_lines.append(line)
                    i += 1
            else:
                formatted_lines.append(line)
                i += 1
        
        return '\n'.join(formatted_lines), changes, has_tables
    
    def _format_table(self, rows: List[str]) -> str:
        """Tablo satırlarını formatla"""
        # Hücreleri ayır
        cells = []
        for row in rows:
            if '\t' in row:
                cells.append([c.strip() for c in row.split('\t')])
            elif '|' in row:
                cells.append([c.strip() for c in row.split('|') if c.strip()])
            else:
                cells.append([c.strip() for c in row.split('  ') if c.strip()])
        
        if not cells:
            return '\n'.join(rows)
        
        # Sütun genişliklerini hesapla
        col_widths = []
        for col in range(len(cells[0])):
            max_width = max(len(str(row[col])) for row in cells if col < len(row))
            col_widths.append(max_width)
        
        # Formatlı tablo oluştur
        formatted_rows = []
        for i, row in enumerate(cells):
            formatted_cells = []
            for j, cell in enumerate(row):
                if j < len(col_widths):
                    formatted_cells.append(cell.ljust(col_widths[j]))
            formatted_rows.append(' | '.join(formatted_cells))
            
            # Başlık satırından sonra ayırıcı ekle
            if i == 0:
                separator = '-+-'.join('-' * w for w in col_widths)
                formatted_rows.append(separator)
        
        return '\n'.join(formatted_rows)
    
    def optimize_sentences(self, text: str, metrics: DocumentMetrics) -> Tuple[str, List[str]]:
        """Cümleleri optimize et"""
        changes = []
        paragraphs = text.split('\n\n')
        optimized_paragraphs = []
        
        for para in paragraphs:
            if not para.strip():
                continue
            
            # Cümlelere böl
            sentences = re.split(r'(?<=[.!?])\s+', para)
            optimized_sentences = []
            
            for sent in sentences:
                if not sent.strip():
                    continue
                
                # Çok uzun cümleleri böl
                words = sent.split()
                if len(words) > self.professional_standards['max_sentence_length']:
                    # Virgül veya bağlaçlardan böl
                    split_points = [i for i, w in enumerate(words) if w.lower() in ['ve', 'ile', 'ancak', 'fakat', 'lakin', 'oysa', 'çünkü']]
                    
                    if split_points:
                        mid = split_points[len(split_points) // 2]
                        sent1 = ' '.join(words[:mid+1])
                        sent2 = ' '.join(words[mid+1:])
                        optimized_sentences.append(sent1 + '.')
                        optimized_sentences.append(sent2[0].upper() + sent2[1:] + '.')
                        changes.append("Uzun cümle bölündü")
                    else:
                        optimized_sentences.append(sent)
                else:
                    optimized_sentences.append(sent)
            
            optimized_paragraphs.append(' '.join(optimized_sentences))
        
        return '\n\n'.join(optimized_paragraphs), changes
    
    def apply_style(self, text: str, style: DocumentStyle) -> str:
        """Stil kurallarını uygula"""
        rules = self.style_rules.get(style, self.style_rules[DocumentStyle.CORPORATE])
        lines = text.split('\n')
        styled_lines = []
        
        for line in lines:
            stripped = line.strip()
            if not stripped:
                styled_lines.append('')
                continue
            
            # Başlık formatı
            if any(re.match(p, stripped) for p in self.professional_standards['heading_patterns']):
                content = re.sub(r'^[#\d.\s]+', '', stripped).strip()
                styled_lines.append(rules['heading_format'].format(content))
            
            # Liste formatı
            elif any(stripped.startswith(i) for i in self.professional_standards['list_indicators']):
                content = stripped.lstrip('•-*→✓✅●○').strip()
                # Liste numarasını bul
                num_match = re.match(r'^(\d+)[.)]', stripped)
                if num_match:
                    styled_lines.append(rules['list_format'].format(num_match.group(1), content))
                else:
                    styled_lines.append(rules['list_format'].format('•', content))
            
            # Normal metin
            else:
                styled_lines.append(line)
        
        return '\n'.join(styled_lines)
    
    # ========== WORD DÜZENLEME (GELİŞTİRİLMİŞ) ==========
    
    def edit_word(self, file_path: str, target_format: str, analysis: Any) -> EditResult:
        """Word dosyası düzenleme - Gelişmiş"""
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor, Cm
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.style import WD_STYLE_TYPE
            
            doc = Document(file_path)
            changes = []
            
            # Metin içeriğini çıkar
            full_text = '\n'.join([p.text for p in doc.paragraphs if p.text.strip()])
            metrics = self.analyze_text_metrics(full_text)
            
            # Stil tespiti
            style = self.detect_style(full_text, metrics)
            rules = self.style_rules.get(style, self.style_rules[DocumentStyle.CORPORATE])
            
            # 1. Sayfa yapısını ayarla
            section = doc.sections[0]
            section.top_margin = Cm(2.5)
            section.bottom_margin = Cm(2.5)
            section.left_margin = Cm(3.0)
            section.right_margin = Cm(3.0)
            changes.append("Sayfa kenar boşlukları ayarlandı")
            
            # 2. Boş paragrafları temizle
            empty_para_count = 0
            paragraphs_to_remove = []
            for i, para in enumerate(doc.paragraphs):
                if not para.text.strip():
                    empty_para_count += 1
                    paragraphs_to_remove.append(i)
            
            for i in reversed(paragraphs_to_remove):
                p = doc.paragraphs[i]._element
                p.getparent().remove(p)
            
            if empty_para_count > 0:
                changes.append(f"{empty_para_count} boş paragraf temizlendi")
            
            # 3. Normal stil ayarları
            normal_style = doc.styles['Normal']
            normal_style.font.name = rules['font_family']
            normal_style.font.size = Pt(rules['font_size'])
            normal_style.paragraph_format.line_spacing = rules['line_spacing']
            normal_style.paragraph_format.space_after = Pt(rules['paragraph_spacing'] * 10)
            
            # 4. Başlık stilleri
            heading_styles = {
                1: (rules['font_family'], rules['font_size'] + 6, True),
                2: (rules['font_family'], rules['font_size'] + 4, True),
                3: (rules['font_family'], rules['font_size'] + 2, True),
            }
            
            for level, (font, size, bold) in heading_styles.items():
                style_name = f'Heading {level}'
                if style_name in doc.styles:
                    h_style = doc.styles[style_name]
                    h_style.font.name = font
                    h_style.font.size = Pt(size)
                    h_style.font.bold = bold
            
            # 5. Başlık tespiti ve atama
            heading_count = 0
            for para in doc.paragraphs:
                text = para.text.strip()
                if text and len(text) < 100:
                    if (text.isupper() or 
                        len(text.split()) <= 5 and not text.endswith(('.', '!', '?')) or
                        any(re.match(p, text) for p in self.professional_standards['heading_patterns'])):
                        
                        # Başlık seviyesini belirle
                        if '#' in text:
                            level = text.count('#')
                            text = text.replace('#', '').strip()
                        elif text.isupper():
                            level = 1
                        else:
                            level = 2
                        
                        para.style = doc.styles[f'Heading {level}']
                        heading_count += 1
                        changes.append(f"Başlık algılandı (Seviye {level}): {text[:30]}...")
            
            # 6. Liste düzenleme
            bullet_count = 0
            number_count = 0
            for para in doc.paragraphs:
                text = para.text.strip()
                if any(text.startswith(i) for i in self.professional_standards['list_indicators']):
                    para.style = doc.styles['List Bullet']
                    para.text = text.lstrip('•-*→✓✅●○').strip()
                    bullet_count += 1
                elif re.match(r'^\d+[.)]', text):
                    para.style = doc.styles['List Number']
                    num = re.match(r'^\d+', text).group()
                    para.text = text.replace(f"{num}.", "").replace(f"{num})", "").strip()
                    number_count += 1
            
            if bullet_count > 0:
                changes.append(f"{bullet_count} madde işaretli liste öğesi düzenlendi")
            if number_count > 0:
                changes.append(f"{number_count} numaralı liste öğesi düzenlendi")
            
            # 7. Tablo düzenlemeleri
            table_count = len(doc.tables)
            if table_count > 0:
                for table in doc.tables:
                    # Tablo stili
                    table.style = 'Light Grid Accent 1'
                    
                    # Başlık satırı
                    if len(table.rows) > 0:
                        for cell in table.rows[0].cells:
                            for para in cell.paragraphs:
                                para.style = doc.styles['Table Grid']
                                for run in para.runs:
                                    run.font.bold = True
                                    run.font.size = Pt(12)
                
                changes.append(f"{table_count} tablo düzenlendi")
            
            # 8. Kalite skoru hesapla
            quality_score = 70
            quality_score += heading_count * 3
            quality_score += bullet_count * 2
            quality_score += number_count * 2
            quality_score += table_count * 5
            quality_score -= empty_para_count * 2
            quality_score = min(100, max(0, quality_score))
            
            # Geçici dosyaya kaydet
            temp_path = file_path.replace('.docx', '_edited.docx').replace('.doc', '_edited.docx')
            doc.save(temp_path)
            
            # Öneriler
            suggestions = []
            if quality_score < 70:
                suggestions.append("Daha fazla başlık ekleyerek belge yapısını iyileştirebilirsiniz")
            if table_count == 0 and metrics.has_tables:
                suggestions.append("Verileriniz için tablo kullanmanız önerilir")
            if metrics.complexity_score > 70:
                suggestions.append("Cümlelerinizi kısaltarak okunabilirliği artırabilirsiniz")
            
            return EditResult(
                success=True,
                edited_content=temp_path,
                changes_made=changes,
                intensity=EditIntensity.PROFESSIONAL,
                processing_time=0,
                word_count=metrics.total_words,
                paragraph_count=metrics.total_paragraphs,
                quality_score=quality_score,
                metrics=metrics,
                style=style,
                suggestions=suggestions
            )
            
        except Exception as e:
            logger.error(f"❌ Word düzenleme hatası: {e}")
            return EditResult(
                success=False,
                edited_content=file_path,
                changes_made=[f"Hata: {str(e)}"],
                intensity=EditIntensity.LIGHT,
                processing_time=0,
                word_count=0,
                paragraph_count=0,
                quality_score=0
            )
    
    # ========== METİN DÜZENLEME (GELİŞTİRİLMİŞ) ==========
    
    def edit_text(self, file_path: str, target_format: str, analysis: Any) -> EditResult:
        """Düz metin dosyası düzenleme - Gelişmiş"""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read()
            
            changes = []
            edited_content = content
            
            # Metrik analizi
            metrics = self.analyze_text_metrics(edited_content)
            
            # Stil tespiti
            style = self.detect_style(edited_content, metrics)
            
            # 1. Metin temizleme
            edited_content, clean_changes = self.clean_text_advanced(edited_content, metrics)
            changes.extend(clean_changes)
            
            # 2. Paragraf formatlama
            edited_content, para_changes = self.format_paragraphs(edited_content, metrics)
            changes.extend(para_changes)
            
            # 3. Liste tespiti ve formatlama
            edited_content, list_changes = self.detect_and_format_lists(edited_content)
            changes.extend(list_changes)
            
            # 4. Tablo tespiti ve formatlama
            edited_content, table_changes, has_tables = self.detect_and_format_tables(edited_content)
            changes.extend(table_changes)
            metrics.has_tables = has_tables
            
            # 5. Cümle optimizasyonu
            if metrics.complexity_score > 60:
                edited_content, sentence_changes = self.optimize_sentences(edited_content, metrics)
                changes.extend(sentence_changes)
            
            # 6. Stil uygulama
            if style != DocumentStyle.CORPORATE:
                edited_content = self.apply_style(edited_content, style)
                changes.append(f"{style.value} stili uygulandı")
            
            # Geçici dosyaya kaydet
            temp_path = file_path.replace('.txt', '_edited.txt')
            with open(temp_path, 'w', encoding='utf-8') as f:
                f.write(edited_content)
            
            # Kalite skoru
            quality_score = 70
            quality_score += len(changes) * 2
            quality_score += metrics.style_consistency
            quality_score = min(100, quality_score)
            
            # Öneriler
            suggestions = []
            if metrics.reading_time_minutes > 10:
                suggestions.append("Belgeniz çok uzun, bölümlere ayırmayı düşünebilirsiniz")
            if metrics.complexity_score > 70:
                suggestions.append("Daha sade bir dil kullanmanız önerilir")
            if not metrics.has_headings and metrics.total_paragraphs > 5:
                suggestions.append("Başlıklar ekleyerek okunabilirliği artırabilirsiniz")
            
            return EditResult(
                success=True,
                edited_content=temp_path,
                changes_made=changes,
                intensity=EditIntensity.PROFESSIONAL,
                processing_time=0,
                word_count=metrics.total_words,
                paragraph_count=metrics.total_paragraphs,
                quality_score=quality_score,
                metrics=metrics,
                style=style,
                suggestions=suggestions
            )
            
        except Exception as e:
            logger.error(f"❌ Metin düzenleme hatası: {e}")
            return EditResult(
                success=False,
                edited_content=file_path,
                changes_made=[f"Hata: {str(e)}"],
                intensity=EditIntensity.LIGHT,
                processing_time=0,
                word_count=0,
                paragraph_count=0,
                quality_score=0
            )
    
    # ========== PDF DÜZENLEME (GELİŞTİRİLMİŞ) ==========
    
    def edit_pdf(self, file_path: str, target_format: str, analysis: Any) -> EditResult:
        """PDF dosyası düzenleme - Gelişmiş"""
        try:
            import PyPDF2
            from reportlab.pdfgen import canvas
            from reportlab.lib.pagesizes import A4
            from reportlab.lib.units import cm
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.platypus import Paragraph, Spacer, SimpleDocTemplate
            from reportlab.lib.enums import TA_JUSTIFY, TA_LEFT, TA_CENTER
            
            changes = []
            
            # PDF'den metin çıkar
            text_content = []
            with open(file_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                total_pages = len(pdf_reader.pages)
                
                for page_num in range(total_pages):
                    page = pdf_reader.pages[page_num]
                    text = page.extract_text()
                    if text.strip():
                        text_content.append(f"[Sayfa {page_num + 1}/{total_pages}]\n{text}")
                    else:
                        changes.append(f"Sayfa {page_num+1} metin içermiyor (OCR gerekebilir)")
            
            # Metin içeriğini birleştir
            full_text = '\n\n'.join(text_content)
            
            # Metrik analizi
            metrics = self.analyze_text_metrics(full_text)
            
            # Metin düzenleme
            text_editor_result = self.edit_text(file_path.replace('.pdf', '.txt'), 'txt', analysis)
            
            if text_editor_result.success:
                with open(text_editor_result.edited_content, 'r', encoding='utf-8') as f:
                    edited_text = f.read()
                
                # Yeni PDF oluştur (profesyonel)
                temp_pdf_path = file_path.replace('.pdf', '_edited.pdf')
                
                # PDF oluşturucu
                doc = SimpleDocTemplate(temp_pdf_path, pagesize=A4,
                                        leftMargin=2*cm, rightMargin=2*cm,
                                        topMargin=2*cm, bottomMargin=2*cm)
                
                styles = getSampleStyleSheet()
                story = []
                
                # Başlık
                title_style = ParagraphStyle(
                    'CustomTitle',
                    parent=styles['Heading1'],
                    fontSize=18,
                    spaceAfter=30,
                    alignment=TA_CENTER,
                    textColor='#2E75B6'
                )
                
                story.append(Paragraph("Düzenlenmiş PDF Belgesi", title_style))
                story.append(Spacer(1, 20))
                
                # Tarih
                date_style = ParagraphStyle(
                    'DateStyle',
                    parent=styles['Normal'],
                    fontSize=10,
                    textColor='#666666',
                    alignment=TA_CENTER
                )
                story.append(Paragraph(f"Oluşturma: {datetime.datetime.now().strftime('%d.%m.%Y %H:%M')}", date_style))
                story.append(Spacer(1, 30))
                
                # Metin içeriği
                normal_style = ParagraphStyle(
                    'CustomNormal',
                    parent=styles['Normal'],
                    fontSize=11,
                    leading=14,
                    alignment=TA_JUSTIFY,
                    spaceAfter=12
                )
                
                paragraphs = edited_text.split('\n\n')
                for para in paragraphs:
                    if para.strip():
                        if para.startswith('[') and 'Sayfa' in para:
                            # Sayfa başlığı
                            p = Paragraph(para, styles['Heading2'])
                            story.append(p)
                        else:
                            p = Paragraph(para.replace('\n', ' '), normal_style)
                            story.append(p)
                        story.append(Spacer(1, 6))
                
                # PDF'i oluştur
                doc.build(story)
                
                changes.extend(text_editor_result.changes_made)
                changes.append(f"PDF yeniden oluşturuldu ({total_pages} sayfa)")
                
                # Geçici dosyaları temizle
                try:
                    os.remove(text_editor_result.edited_content)
                except:
                    pass
                
                return EditResult(
                    success=True,
                    edited_content=temp_pdf_path,
                    changes_made=changes,
                    intensity=EditIntensity.HEAVY,
                    processing_time=0,
                    word_count=metrics.total_words,
                    paragraph_count=metrics.total_paragraphs,
                    quality_score=text_editor_result.quality_score,
                    metrics=metrics,
                    style=text_editor_result.style,
                    suggestions=text_editor_result.suggestions
                )
            else:
                return text_editor_result
                
        except Exception as e:
            logger.error(f"❌ PDF düzenleme hatası: {e}")
            return EditResult(
                success=False,
                edited_content=file_path,
                changes_made=[f"Hata: {str(e)}"],
                intensity=EditIntensity.LIGHT,
                processing_time=0,
                word_count=0,
                paragraph_count=0,
                quality_score=0
            )
    
    # ========== MARKDOWN DÜZENLEME ==========
    
    def edit_markdown(self, file_path: str, target_format: str, analysis: Any) -> EditResult:
        """Markdown dosyası düzenleme"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            changes = []
            
            # Metrik analizi
            metrics = self.analyze_text_metrics(content)
            
            # Stil tespiti
            style = self.detect_style(content, metrics)
            
            # Markdown düzenleme
            lines = content.split('\n')
            formatted_lines = []
            
            in_code_block = False
            in_table = False
            
            for line in lines:
                # Kod bloğu kontrolü
                if line.startswith('```'):
                    in_code_block = not in_code_block
                    formatted_lines.append(line)
                    continue
                
                if in_code_block:
                    formatted_lines.append(line)
                    continue
                
                # Tablo kontrolü
                if '|' in line and '-' in line and (len(line.split('|')) > 2):
                    in_table = not in_table
                    formatted_lines.append(line)
                    continue
                
                if in_table:
                    formatted_lines.append(line)
                    continue
                
                # Başlık düzenleme
                if line.startswith('#'):
                    level = line.count('#')
                    content = line.replace('#', '').strip()
                    formatted_lines.append(f"{'#' * level} {content}")
                    if level == 1:
                        changes.append(f"Ana başlık: {content[:30]}...")
                
                # Liste düzenleme
                elif line.strip().startswith(('-', '*', '+')):
                    content = line.strip()[1:].strip()
                    formatted_lines.append(f"- {content}")
                
                # Numaralı liste
                elif re.match(r'^\d+\.', line.strip()):
                    num = re.match(r'^\d+', line.strip()).group()
                    content = re.sub(r'^\d+\.', '', line.strip()).strip()
                    formatted_lines.append(f"{num}. {content}")
                
                # Normal metin
                else:
                    # Çok uzun satırları böl
                    if len(line) > 80:
                        words = line.split()
                        current_line = []
                        current_length = 0
                        
                        for word in words:
                            if current_length + len(word) + 1 <= 80:
                                current_line.append(word)
                                current_length += len(word) + 1
                            else:
                                if current_line:
                                    formatted_lines.append(' '.join(current_line))
                                current_line = [word]
                                current_length = len(word)
                        
                        if current_line:
                            formatted_lines.append(' '.join(current_line))
                    else:
                        formatted_lines.append(line)
            
            edited_content = '\n'.join(formatted_lines)
            
            # Geçici dosyaya kaydet
            temp_path = file_path.replace('.md', '_edited.md')
            with open(temp_path, 'w', encoding='utf-8') as f:
                f.write(edited_content)
            
            # Kalite skoru
            quality_score = 80 + len(changes) * 2
            quality_score = min(100, quality_score)
            
            return EditResult(
                success=True,
                edited_content=temp_path,
                changes_made=changes,
                intensity=EditIntensity.MEDIUM,
                processing_time=0,
                word_count=metrics.total_words,
                paragraph_count=metrics.total_paragraphs,
                quality_score=quality_score,
                metrics=metrics,
                style=style
            )
            
        except Exception as e:
            logger.error(f"❌ Markdown düzenleme hatası: {e}")
            return EditResult(
                success=False,
                edited_content=file_path,
                changes_made=[f"Hata: {str(e)}"],
                intensity=EditIntensity.LIGHT,
                processing_time=0,
                word_count=0,
                paragraph_count=0,
                quality_score=0
            )
    
    # ========== GÖRSEL DÜZENLEME (GELİŞTİRİLMİŞ) ==========
    
    def edit_image(self, file_path: str, target_format: str, analysis: Any) -> EditResult:
        """Görsel dosyası düzenleme - Gelişmiş"""
        try:
            from PIL import Image, ImageEnhance, ImageFilter, ImageOps
            import pytesseract
            
            changes = []
            
            # 1. Görseli yükle ve analiz et
            image = Image.open(file_path)
            original_format = image.format
            original_mode = image.mode
            original_size = image.size
            
            changes.append(f"Orijinal görsel: {original_size[0]}x{original_size[1]}, {original_mode}")
            
            # 2. Görsel ön işleme
            # Görseli büyüt (OCR kalitesi için)
            width, height = image.size
            if width < 2000:
                new_size = (width * 2, height * 2)
                image = image.resize(new_size, Image.Resampling.LANCZOS)
                changes.append(f"Görsel büyütüldü: {width}x{height} -> {new_size[0]}x{new_size[1]}")
            
            # Gri tonlamaya çevir
            if image.mode != 'L':
                image = image.convert('L')
                changes.append("Gri tonlamaya çevrildi")
            
            # Kontrast artır
            enhancer = ImageEnhance.Contrast(image)
            image = enhancer.enhance(2.5)
            changes.append("Kontrast %150 artırıldı")
            
            # Gürültü azalt
            image = image.filter(ImageFilter.MedianFilter(size=3))
            image = image.filter(ImageFilter.SHARPEN)
            changes.append("Gürültü azaltıldı, keskinleştirildi")
            
            # Kenar iyileştirme
            image = image.filter(ImageFilter.EDGE_ENHANCE)
            changes.append("Kenar iyileştirmesi uygulandı")
            
            # 3. Geçici görsel kaydet
            temp_image_path = file_path + "_enhanced.png"
            image.save(temp_image_path, 'PNG', dpi=(300,300))
            
            # 4. OCR uygula (çoklu deneme)
            text = ""
            ocr_configs = [
                r'--oem 3 --psm 6 -l tur+eng',   # Varsayılan
                r'--oem 3 --psm 3 -l tur+eng',   # Otomatik sayfa bölme
                r'--oem 3 --psm 4 -l tur+eng',   # Tek sütun
                r'--oem 3 --psm 11 -l tur+eng',  # Sparse text
            ]
            
            for config in ocr_configs:
                try:
                    result = pytesseract.image_to_string(temp_image_path, config=config)
                    if result.strip() and len(result) > len(text):
                        text = result
                except:
                    continue
            
            # 5. OCR METNİNİ TEMİZLE VE DÜZENLE
            if text.strip():
                # Yaygın OCR hatalarını düzelt
                text = fix_common_ocr_errors(text)
                
                # Gereksiz satırları filtrele
                cleaned_text, removed_lines = clean_ocr_text(text)
                
                # Fazla boşlukları temizle
                cleaned_text = re.sub(r'\s+', ' ', cleaned_text)
                cleaned_text = re.sub(r'\n\s*\n\s*\n', '\n\n', cleaned_text)
                
                # Anlamlı satırları birleştir
                cleaned_text = merge_intelligent_lines(cleaned_text)
                
                # Normalize et
                cleaned_text = normalize_whitespace(cleaned_text)
                
                # OCR güven skoru hesapla
                ocr_confidence = calculate_ocr_confidence(cleaned_text)
                
                changes.append(f"OCR başarılı ({len(text)} karakter, {len(removed_lines)} gereksiz satır temizlendi)")
                changes.append(f"OCR güven skoru: %{ocr_confidence:.1f}")
                
                text = cleaned_text
            else:
                changes.append("OCR başarısız, görselde metin olmayabilir")
                ocr_confidence = 0.0
            
            # 6. OCR çıktısını metin olarak kaydet
            temp_txt_path = file_path + "_ocr.txt"
            with open(temp_txt_path, 'w', encoding='utf-8') as f:
                f.write(text)
            
            # 7. Metin düzenleyiciyi çağır
            text_editor = AIEditor()
            text_result = text_editor.edit_text(temp_txt_path, 'txt', analysis)
            
            if text_result.success:
                changes.extend(text_result.changes_made)
                
                # 8. Hedef formata göre çıktı hazırla
                metrics = self.analyze_text_metrics(text)
                metrics.ocr_confidence = ocr_confidence
                
                if target_format == 'WORD':
                    from docx import Document
                    from docx.shared import Inches, Pt
                    
                    doc = Document()
                    
                    # Başlık
                    title = doc.add_heading('Görselden OCR ile Dönüştürülen Metin', 0)
                    title.alignment = 1  # Center
                    
                    # Kaynak bilgisi
                    doc.add_paragraph(f"Kaynak dosya: {os.path.basename(file_path)}")
                    doc.add_paragraph(f"Görsel boyutu: {original_size[0]}x{original_size[1]}")
                    doc.add_paragraph(f"Dönüşüm tarihi: {datetime.datetime.now().strftime('%d.%m.%Y %H:%M')}")
                    doc.add_paragraph(f"OCR güven skoru: %{ocr_confidence:.1f}")
                    doc.add_paragraph()
                    
                    # Düzenlenmiş metin
                    with open(text_result.edited_content, 'r', encoding='utf-8') as f:
                        edited_text = f.read()
                    
                    for para in edited_text.split('\n\n'):
                        if para.strip():
                            doc.add_paragraph(para.strip())
                    
                    # Orijinal görsel
                    doc.add_page_break()
                    doc.add_heading('Orijinal Görsel', level=1)
                    doc.add_picture(file_path, width=Inches(5))
                    
                    output_path = file_path.replace('.', '_edited.') + '.docx'
                    doc.save(output_path)
                    
                elif target_format == 'PDF':
                    from reportlab.pdfgen import canvas
                    from reportlab.lib.pagesizes import A4
                    from reportlab.lib.units import cm
                    
                    c = canvas.Canvas(file_path.replace('.', '_edited.') + '.pdf', pagesize=A4)
                    width, height = A4
                    
                    y = height - 2*cm
                    line_height = 0.5*cm
                    
                    with open(text_result.edited_content, 'r', encoding='utf-8') as f:
                        edited_text = f.read()
                    
                    # Başlık
                    c.setFont("Helvetica-Bold", 16)
                    c.drawString(2*cm, y, "Görselden OCR ile Dönüştürülen Metin")
                    y -= line_height * 2
                    
                    # Kaynak bilgisi
                    c.setFont("Helvetica", 10)
                    c.drawString(2*cm, y, f"Kaynak: {os.path.basename(file_path)}")
                    y -= line_height
                    c.drawString(2*cm, y, f"OCR Güven: %{ocr_confidence:.1f}")
                    y -= line_height * 2
                    
                    # İçerik
                    c.setFont("Helvetica", 11)
                    for line in edited_text.split('\n'):
                        if y < line_height + 1*cm:
                            c.showPage()
                            y = height - 2*cm
                            c.setFont("Helvetica", 11)
                        
                        if line.strip():
                            # Satırları kısalt
                            if len(line) > 80:
                                line = line[:80] + "..."
                            c.drawString(2*cm, y, line)
                        y -= line_height
                    
                    c.save()
                    output_path = file_path.replace('.', '_edited.') + '.pdf'
                else:
                    output_path = text_result.edited_content
                
                # Geçici dosyaları temizle
                try:
                    os.remove(temp_image_path)
                    os.remove(temp_txt_path)
                except:
                    pass
                
                # Kalite skoru
                quality_score = 70
                if text.strip():
                    quality_score += 20
                if len(text) > 100:
                    quality_score += 10
                if ocr_confidence > 80:
                    quality_score += 10
                elif ocr_confidence > 50:
                    quality_score += 5
                quality_score = min(100, quality_score)
                
                return EditResult(
                    success=True,
                    edited_content=output_path,
                    changes_made=changes,
                    intensity=EditIntensity.HEAVY,
                    processing_time=0,
                    word_count=metrics.total_words,
                    paragraph_count=metrics.total_paragraphs,
                    quality_score=quality_score,
                    metrics=metrics
                )
            else:
                return text_result
                
        except Exception as e:
            logger.error(f"❌ Görsel düzenleme hatası: {e}")
            import traceback
            traceback.print_exc()
            return EditResult(
                success=False,
                edited_content=file_path,
                changes_made=[f"Hata: {str(e)}"],
                intensity=EditIntensity.LIGHT,
                processing_time=0,
                word_count=0,
                paragraph_count=0,
                quality_score=0
            )
    
    # ========== ANA DÜZENLEME FONKSİYONU ==========
    
    def edit(self, file_path: str, target_format: str, analysis_result: Any = None, 
             style: Optional[DocumentStyle] = None) -> EditResult:
        """
        Ana düzenleme fonksiyonu - Gelişmiş
        Dosya türüne göre uygun editörü çağırır
        """
        import time
        start_time = time.time()
        
        try:
            file_ext = os.path.splitext(file_path)[1].lower().replace('.', '')
            
            if file_ext not in self.supported_formats:
                logger.warning(f"⚠️ Desteklenmeyen format: {file_ext}, direkt dönüşüm yapılacak")
                return EditResult(
                    success=True,
                    edited_content=file_path,
                    changes_made=["Desteklenmeyen format, düzenleme yapılmadı"],
                    intensity=EditIntensity.LIGHT,
                    processing_time=time.time() - start_time,
                    word_count=0,
                    paragraph_count=0,
                    quality_score=50
                )
            
            # Uygun editörü çağır
            editor = self.supported_formats[file_ext]
            result = editor(file_path, target_format, analysis_result)
            
            # İstenen stili uygula
            if style and result.success and result.edited_content:
                if file_ext in ['docx', 'doc']:
                    # Word için stil uygula
                    self._apply_style_to_word(result.edited_content, style)
                elif file_ext in ['txt', 'md']:
                    # Metin için stil uygula
                    with open(result.edited_content, 'r', encoding='utf-8') as f:
                        content = f.read()
                    styled_content = self.apply_style(content, style)
                    with open(result.edited_content, 'w', encoding='utf-8') as f:
                        f.write(styled_content)
            
            result.processing_time = time.time() - start_time
            logger.info(f"✅ Düzenleme tamamlandı: {file_path} -> {target_format} | Süre: {result.processing_time:.2f}s")
            return result
            
        except Exception as e:
            logger.error(f"❌ Düzenleme hatası: {e}")
            import traceback
            traceback.print_exc()
            return EditResult(
                success=False,
                edited_content=None,
                changes_made=[f"Hata: {str(e)}"],
                intensity=EditIntensity.LIGHT,
                processing_time=time.time() - start_time,
                word_count=0,
                paragraph_count=0,
                quality_score=0
            )
    
    def _apply_style_to_word(self, doc_path: str, style: DocumentStyle):
        """Word belgesine stil uygula"""
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor, Cm
            
            doc = Document(doc_path)
            rules = self.style_rules.get(style, self.style_rules[DocumentStyle.CORPORATE])
            
            # Normal stil
            normal_style = doc.styles['Normal']
            normal_style.font.name = rules['font_family']
            normal_style.font.size = Pt(rules['font_size'])
            normal_style.paragraph_format.line_spacing = rules['line_spacing']
            
            doc.save(doc_path)
            
        except Exception as e:
            logger.error(f"❌ Stil uygulama hatası: {e}")
    
    # ========== DÜZENLEME ÖZETİ ==========
    
    def get_edit_summary(self, result: EditResult) -> str:
        """Gelişmiş düzenleme özeti oluştur"""
        summary = f"""
📝 **AKILLI DÜZENLEME RAPORU**

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
📋 **BELGE BİLGİLERİ**
• Stil: {result.style.value.upper()}
• Yoğunluk: {result.intensity.value}
• Kelime Sayısı: {result.word_count:,}
• Paragraf Sayısı: {result.paragraph_count}
• Kalite Puanı: %{result.quality_score}

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
✨ **YAPILAN İŞLEMLER:**
"""
        for change in result.changes_made:
            summary += f"  • {change}\n"
        
        if result.metrics:
            summary += f"""
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
📊 **METRİKLER**
• Okuma Süresi: {result.metrics.reading_time_minutes:.1f} dakika
• Ortalama Cümle: {result.metrics.avg_sentence_length:.1f} kelime
• Karmaşıklık: %{result.metrics.complexity_score}
• Dil: {result.metrics.language.upper()}
"""
            if hasattr(result.metrics, 'ocr_confidence') and result.metrics.ocr_confidence > 0:
                summary += f"• OCR Güven: %{result.metrics.ocr_confidence:.1f}\n"
        
        if result.suggestions:
            summary += """
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
💡 **İYİLEŞTİRME ÖNERİLERİ**
"""
            for suggestion in result.suggestions:
                summary += f"  • {suggestion}\n"
        
        summary += """
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
✅ **DOSYA PROFESYONEL BELGE STANDARDINA UYGUN HALE GETİRİLDİ.**
"""
        return summary


# ========== KULLANIM KOLAYLIĞI FONKSİYONLARI ==========

def smart_edit(file_path: str, target_format: str, analysis_result: Any = None, 
               style: str = None) -> Tuple[bool, str, List[str], str]:
    """
    Akıllı düzenleme yap
    Returns: (başarılı_mı, çıktı_dosyası, değişiklikler, özet)
    """
    editor = AIEditor()
    
    style_map = {
        'kurumsal': DocumentStyle.CORPORATE,
        'resmi': DocumentStyle.OFFICIAL,
        'akademik': DocumentStyle.ACADEMIC,
        'teknik': DocumentStyle.TECHNICAL,
        'hukuki': DocumentStyle.LEGAL,
        'finansal': DocumentStyle.FINANCIAL
    }
    
    style_enum = style_map.get(style, None)
    result = editor.edit(file_path, target_format, analysis_result, style_enum)
    
    if result.success:
        summary = editor.get_edit_summary(result)
        return True, result.edited_content, result.changes_made, summary
    else:
        return False, file_path, result.changes_made, "Düzenleme başarısız oldu."

def analyze_document(file_path: str, file_type: str) -> Dict:
    """
    Belgeyi analiz et ve metrikleri döndür
    """
    editor = AIEditor()
    
    # Metin çıkar
    text = ""
    try:
        if file_type == 'WORD':
            from docx import Document
            doc = Document(file_path)
            text = '\n'.join([p.text for p in doc.paragraphs if p.text.strip()])
        elif file_type == 'PDF':
            import PyPDF2
            with open(file_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
        elif file_type == 'EXCEL':
            import pandas as pd
            df = pd.read_excel(file_path)
            text = df.to_string()
        elif file_type == 'TXT':
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                text = f.read()
        else:
            return {"error": "Desteklenmeyen format"}
    except Exception as e:
        return {"error": str(e)}
    
    metrics = editor.analyze_text_metrics(text)
    style = editor.detect_style(text, metrics)
    
    return {
        'metrics': {
            'total_words': metrics.total_words,
            'unique_words': metrics.unique_words,
            'avg_word_length': metrics.avg_word_length,
            'total_sentences': metrics.total_sentences,
            'avg_sentence_length': metrics.avg_sentence_length,
            'total_paragraphs': metrics.total_paragraphs,
            'avg_paragraph_length': metrics.avg_paragraph_length,
            'reading_time_minutes': metrics.reading_time_minutes,
            'complexity_score': metrics.complexity_score,
            'language': metrics.language
        },
        'style': style.value,
        'suggestions': []
    }


# ========== TEST FONKSİYONU ==========
if __name__ == "__main__":
    print("🔧 Akıllı Düzenleme Modülü Test Ediliyor...")
    print("=" * 60)
    
    editor = AIEditor()
    
    # Test metni
    test_text = """
    BU BİR TEST BAŞLIĞIDIR
    
    Bu bir test paragrafıdır. İçinde birden fazla cümle bulunur. 
    Bu cümleler otomatik olarak düzenlenecek ve optimize edilecektir.
    
    • Madde 1
    • Madde 2
    • Madde 3
    
    1. Numaralı madde
    2. Numaralı madde
    3. Numaralı madde
    """
    
    # Test
    metrics = editor.analyze_text_metrics(test_text)
    style = editor.detect_style(test_text, metrics)
    
    print(f"📊 Metrikler:")
    print(f"  • Kelime Sayısı: {metrics.total_words}")
    print(f"  • Cümle Sayısı: {metrics.total_sentences}")
    print(f"  • Paragraf Sayısı: {metrics.total_paragraphs}")
    print(f"  • Okuma Süresi: {metrics.reading_time_minutes:.1f} dakika")
    print(f"  • Karmaşıklık: %{metrics.complexity_score}")
    print(f"  • Tespit Edilen Stil: {style.value}")
    print("=" * 60)
    
    # Örnek düzenleme raporu
    test_result = EditResult(
        success=True,
        edited_content="test_edited.docx",
        changes_made=[
            "5 boş paragraf temizlendi",
            "3 başlık algılandı ve formatlandı",
            "2 liste öğesi düzenlendi",
            "1 tablo düzenlendi",
            "Kurumsal stil uygulandı"
        ],
        intensity=EditIntensity.PROFESSIONAL,
        processing_time=2.5,
        word_count=1250,
        paragraph_count=48,
        quality_score=94,
        metrics=metrics,
        style=DocumentStyle.CORPORATE,
        suggestions=[
            "Daha fazla başlık ekleyerek belge yapısını iyileştirebilirsiniz",
            "Cümlelerinizi kısaltarak okunabilirliği artırabilirsiniz"
        ]
    )
    
    print(editor.get_edit_summary(test_result))