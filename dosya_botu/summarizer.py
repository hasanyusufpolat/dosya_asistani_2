"""
PROFESYONEL BELGE ÖZETLEME MODÜLÜ
Tüm dosya türleri için akıllı ve anlamlı özet çıkarır
Gelişmiş metin analizi, anahtar kelime çıkarımı ve çoklu dil desteği
"""

import re
import logging
import math
from typing import Dict, List, Optional, Tuple, Set
from dataclasses import dataclass, field
from enum import Enum
from collections import Counter

# Loglama ayarları
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('summarizer.log'),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

class SummaryLength(Enum):
    """Özet uzunluk seviyeleri"""
    SHORT = "kısa"        # 3-4 satır
    MEDIUM = "orta"        # 5-7 satır
    DETAILED = "detaylı"   # 8-10 satır

class SummaryStyle(Enum):
    """Özet stilleri"""
    BULLET = "madde"           # Madde işaretli
    PARAGRAPH = "paragraf"     # Paragraf formatında
    STRUCTURED = "yapısal"     # Başlıklı yapısal

@dataclass
class SummaryMetrics:
    """Özet metrikleri"""
    original_length: int = 0
    summary_length: int = 0
    compression_ratio: float = 0.0
    keyword_count: int = 0
    sentence_count: int = 0
    important_terms: List[str] = field(default_factory=list)
    language: str = "unknown"
    readability_score: float = 0.0

@dataclass
class SummaryResult:
    """Özetleme sonucu (gelişmiş)"""
    summary: str
    key_points: List[str]
    word_count: int
    confidence: int
    document_type: str
    metrics: SummaryMetrics
    style: SummaryStyle
    length: SummaryLength
    warnings: List[str] = field(default_factory=list)
    suggestions: List[str] = field(default_factory=list)

class DocumentSummarizer:
    """Profesyonel Belge Özetleyici - Gelişmiş Versiyon"""
    
    def __init__(self, length: SummaryLength = SummaryLength.MEDIUM, 
                 style: SummaryStyle = SummaryStyle.BULLET):
        self.length = length
        self.style = style
        
        # Önemli bilgi desenleri (genişletilmiş)
        self.important_patterns = {
            'belge_turu': {
                'tr': [r'fatura', r'dekont', r'bordro', r'beyanname', r'rapor', r'tablo', r'ekstre', r'hesap'],
                'en': [r'invoice', r'receipt', r'payroll', r'declaration', r'report', r'table', r'statement', r'account']
            },
            'tarih': {
                'tr': [
                    r'tarih\s*:?\s*(\d{1,2}[\.\/]\d{1,2}[\.\/]\d{2,4})',
                    r'(\d{1,2})[\.\/](\d{1,2})[\.\/](\d{2,4})'
                ],
                'en': [
                    r'date\s*:?\s*(\d{1,2}[\.\/]\d{1,2}[\.\/]\d{2,4})',
                    r'(\d{1,2})[\.\/](\d{1,2})[\.\/](\d{2,4})'
                ]
            },
            'firma': {
                'tr': [
                    r'(?:firma|şirket|company)\s*:?\s*([A-Za-zğüşıöçĞÜŞİÖÇ\s\.]+)',
                    r'(?:alıcı|satıcı|customer|supplier)\s*:?\s*([A-Za-zğüşıöçĞÜŞİÖÇ\s\.]+)'
                ],
                'en': [
                    r'(?:company|firm|corporation)\s*:?\s*([A-Za-z\s\.]+)',
                    r'(?:customer|supplier|buyer|seller)\s*:?\s*([A-Za-z\s\.]+)'
                ]
            },
            'tutar': {
                'tr': [
                    r'toplam\s*:?\s*([\d.,]+\s*(?:tl|try|₺|usd|eur|\$|€))',
                    r'genel\s*toplam\s*:?\s*([\d.,]+\s*(?:tl|try|₺|usd|eur|\$|€))',
                    r'([\d.,]+)\s*(?:tl|try|₺|usd|eur|\$|€)'
                ],
                'en': [
                    r'total\s*:?\s*([\d.,]+\s*(?:usd|eur|\$|€|gbp|jpy))',
                    r'amount\s*:?\s*([\d.,]+\s*(?:usd|eur|\$|€|gbp|jpy))',
                    r'([\d.,]+)\s*(?:usd|eur|\$|€|gbp|jpy)'
                ]
            },
            'kdv': {
                'tr': [
                    r'kdv\s*:?\s*([\d.,]+\s*%)',
                    r'kdv\s*orani\s*:?\s*([\d.,]+)'
                ],
                'en': [
                    r'vat\s*:?\s*([\d.,]+\s*%)',
                    r'vat\s*rate\s*:?\s*([\d.,]+)'
                ]
            },
            'vergi_no': {
                'tr': [
                    r'vergi\s+no\s*:?\s*(\d+)',
                    r'vkn\s*:?\s*(\d+)'
                ],
                'en': [
                    r'tax\s+id\s*:?\s*(\d+)',
                    r'tax\s+number\s*:?\s*(\d+)'
                ]
            },
            'iban': {
                'tr': [
                    r'iban\s*:?\s*([A-Z]{2}\d{2}\s*(?:\d{4}\s*){4,7})',
                    r'iban\s*:?\s*([A-Z]{2}\d{2}[A-Z0-9]{1,30})'
                ],
                'en': [
                    r'iban\s*:?\s*([A-Z]{2}\d{2}\s*(?:\d{4}\s*){4,7})',
                    r'iban\s*:?\s*([A-Z]{2}\d{2}[A-Z0-9]{1,30})'
                ]
            },
            'tel': {
                'tr': [
                    r'tel\s*:?\s*([0-9\s\(\)\+\-]{7,})',
                    r'telefon\s*:?\s*([0-9\s\(\)\+\-]{7,})'
                ],
                'en': [
                    r'phone\s*:?\s*([0-9\s\(\)\+\-]{7,})',
                    r'tel\s*:?\s*([0-9\s\(\)\+\-]{7,})'
                ]
            },
            'email': {
                'tr': [
                    r'email\s*:?\s*([a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,})',
                    r'e-posta\s*:?\s*([a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,})'
                ],
                'en': [
                    r'email\s*:?\s*([a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,})',
                    r'e-mail\s*:?\s*([a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,})'
                ]
            },
            'fatura_no': {
                'tr': [
                    r'fatura\s*no\s*:?\s*([A-Z0-9\-]+)',
                    r'invoice\s*no\s*:?\s*([A-Z0-9\-]+)'
                ],
                'en': [
                    r'invoice\s*no\s*:?\s*([A-Z0-9\-]+)',
                    r'inv\s*#\s*:?\s*([A-Z0-9\-]+)'
                ]
            }
        }
        
        # Kritik notlar için desenler (genişletilmiş)
        self.critical_patterns = [
            (r'geçersiz|invalid', '⚠️ Geçersiz belge'),
            (r'iptal|cancelled|canceled', '❌ İptal edilmiş'),
            (r'vadesi geçmiş|overdue|expired', '⏰ Vadesi geçmiş'),
            (r'eksik|missing|incomplete', '📋 Eksik bilgi'),
            (r'hata|error|mistake', '❌ Hatalı'),
            (r'uyuşmazlık|discrepancy|mismatch', '⚠️ Tutar uyuşmazlığı'),
            (r'tehir|delay|gecikme', '⏱️ Gecikmeli'),
            (r'onay|approved|onaylandı', '✅ Onaylanmış'),
            (r'red|rejected|reddedildi', '❌ Reddedilmiş'),
            (r'beklemede|pending|waiting', '⏳ Beklemede')
        ]
        
        # Dil karakter setleri
        self.language_chars = {
            'tr': set('ğüşıöçĞÜŞİÖÇ'),
            'en': set(),
            'de': set('äöüßÄÖÜ'),
            'fr': set('éèêëàâçôùûÿœæ'),
            'es': set('ñáéíóúüÑÁÉÍÓÚÜ'),
            'it': set('àèéìíîòóùú')
        }
        
        # Özet uzunluk ayarları
        self.length_settings = {
            SummaryLength.SHORT: {'max_lines': 4, 'max_points': 3},
            SummaryLength.MEDIUM: {'max_lines': 7, 'max_points': 5},
            SummaryLength.DETAILED: {'max_lines': 10, 'max_points': 8}
        }
    
    def summarize(self, text: str, document_type: str, 
                 length: Optional[SummaryLength] = None,
                 style: Optional[SummaryStyle] = None) -> SummaryResult:
        """
        Belgeyi özetle (gelişmiş)
        
        Args:
            text: Özetlenecek metin
            document_type: Belge türü
            length: Özet uzunluğu (opsiyonel)
            style: Özet stili (opsiyonel)
        """
        try:
            # Parametreleri ayarla
            length = length or self.length
            style = style or self.style
            len_settings = self.length_settings[length]
            
            # Metrikler
            metrics = SummaryMetrics()
            metrics.original_length = len(text.split())
            metrics.language = self._detect_language(text)
            
            # 1. Önemli bilgileri çıkar
            extracted = self._extract_important_info_advanced(text, metrics.language)
            
            # 2. Anahtar kelimeleri bul
            keywords = self._extract_keywords(text)
            metrics.keyword_count = len(keywords)
            metrics.important_terms = keywords[:5]
            
            # 3. Kritik notları kontrol et
            critical_notes = self._check_critical_notes_advanced(text)
            
            # 4. Özet satırlarını oluştur (stile göre)
            summary_lines = self._build_summary_lines(
                extracted, critical_notes, document_type, style, len_settings
            )
            
            # 5. Özet metnini oluştur
            summary = self._format_summary(summary_lines, style)
            
            # 6. Anahtar noktaları çıkar
            key_points = self._extract_key_points_advanced(
                extracted, critical_notes, keywords, len_settings['max_points']
            )
            
            # 7. Güven skoru hesapla
            confidence, warnings = self._calculate_confidence_advanced(extracted, metrics)
            
            # 8. Öneriler oluştur
            suggestions = self._generate_suggestions(extracted, metrics)
            
            # Metrikleri güncelle
            metrics.summary_length = len(summary.split())
            metrics.sentence_count = len([s for s in summary.split('.') if s.strip()])
            metrics.compression_ratio = metrics.summary_length / metrics.original_length if metrics.original_length > 0 else 0
            metrics.readability_score = self._calculate_readability(summary)
            
            logger.info(f"✅ Belge özeti oluşturuldu - {len(summary_lines)} satır, sıkıştırma: %{metrics.compression_ratio*100:.1f}")
            
            return SummaryResult(
                summary=summary,
                key_points=key_points,
                word_count=metrics.summary_length,
                confidence=confidence,
                document_type=document_type,
                metrics=metrics,
                style=style,
                length=length,
                warnings=warnings,
                suggestions=suggestions
            )
            
        except Exception as e:
            logger.error(f"❌ Özetleme hatası: {e}")
            import traceback
            traceback.print_exc()
            
            metrics = SummaryMetrics()
            return SummaryResult(
                summary="⚠️ Özet oluşturulamadı.",
                key_points=[],
                word_count=0,
                confidence=0,
                document_type=document_type,
                metrics=metrics,
                style=style or self.style,
                length=length or self.length,
                warnings=[str(e)]
            )
    
    def _detect_language(self, text: str) -> str:
        """Metnin dilini tespit et"""
        if not text:
            return 'unknown'
        
        scores = {}
        for lang, chars in self.language_chars.items():
            count = sum(1 for c in text if c in chars)
            if count > 0:
                scores[lang] = count
        
        if scores:
            return max(scores, key=scores.get)
        return 'en'
    
    def _extract_important_info_advanced(self, text: str, language: str) -> Dict:
        """Gelişmiş önemli bilgi çıkarımı"""
        extracted = {}
        
        for field, lang_patterns in self.important_patterns.items():
            # Dile özel patternleri dene
            patterns = lang_patterns.get(language, [])
            # İngilizce'yi de dene
            if language != 'en':
                patterns.extend(lang_patterns.get('en', []))
            
            for pattern in patterns:
                match = re.search(pattern, text, re.IGNORECASE | re.MULTILINE)
                if match:
                    if match.groups():
                        value = match.group(1).strip()
                    else:
                        value = match.group(0).strip()
                    
                    # Değeri temizle
                    value = re.sub(r'\s+', ' ', value)
                    extracted[field] = value
                    break
        
        return extracted
    
    def _extract_keywords(self, text: str, top_n: int = 10) -> List[str]:
        """Anahtar kelimeleri çıkar"""
        # Stop words (Türkçe ve İngilizce)
        stop_words = set([
            've', 'ile', 'için', 'bir', 'bu', 'şu', 'o', 'da', 'de',
            'the', 'and', 'for', 'with', 'this', 'that', 'from',
            'ancak', 'fakat', 'lakin', 'çünkü', 'veya', 'ya da'
        ])
        
        # Kelimeleri temizle
        words = re.findall(r'\b[a-zA-ZğüşıöçĞÜŞİÖÇ]{3,}\b', text.lower())
        
        # Stop words'leri çıkar
        words = [w for w in words if w not in stop_words]
        
        # Frekansları hesapla
        word_freq = Counter(words)
        
        # En sık geçen kelimeleri al
        return [word for word, _ in word_freq.most_common(top_n)]
    
    def _check_critical_notes_advanced(self, text: str) -> List[str]:
        """Gelişmiş kritik not kontrolü"""
        notes = []
        text_lower = text.lower()
        
        for pattern, note in self.critical_patterns:
            if re.search(pattern, text_lower, re.IGNORECASE):
                # Aynı nottan birden fazla varsa tekil olarak ekle
                if note not in notes:
                    notes.append(note)
        
        return notes
    
    def _build_summary_lines(self, extracted: Dict, critical_notes: List[str],
                            document_type: str, style: SummaryStyle,
                            settings: Dict) -> List[str]:
        """Özet satırlarını oluştur"""
        lines = []
        
        # Belge türü (her stilde var)
        doc_type_display = self._format_document_type(document_type)
        if style == SummaryStyle.BULLET:
            lines.append(f"📄 **Belge Türü:** {doc_type_display}")
        elif style == SummaryStyle.STRUCTURED:
            lines.append(f"📄 BELGE TÜRÜ: {doc_type_display}")
            lines.append("─" * 40)
        else:
            lines.append(f"Belge Türü: {doc_type_display}")
        
        # Tarih
        if extracted.get('tarih'):
            lines.append(self._format_line("📅 Tarih", extracted['tarih'], style))
        
        # Firma
        if extracted.get('firma'):
            lines.append(self._format_line("🏢 Firma", extracted['firma'], style))
        
        # Tutar
        if extracted.get('tutar'):
            lines.append(self._format_line("💰 Tutar", extracted['tutar'], style))
        
        # KDV
        if extracted.get('kdv'):
            lines.append(self._format_line("📊 KDV", extracted['kdv'], style))
        
        # Vergi No
        if extracted.get('vergi_no'):
            lines.append(self._format_line("🔢 Vergi No", extracted['vergi_no'], style))
        
        # IBAN
        if extracted.get('iban'):
            iban_short = extracted['iban'][:20] + "..." if len(extracted['iban']) > 20 else extracted['iban']
            lines.append(self._format_line("🏦 IBAN", iban_short, style))
        
        # Telefon
        if extracted.get('tel'):
            lines.append(self._format_line("📞 Telefon", extracted['tel'], style))
        
        # Email
        if extracted.get('email'):
            lines.append(self._format_line("✉️ Email", extracted['email'], style))
        
        # Fatura No
        if extracted.get('fatura_no'):
            lines.append(self._format_line("📋 Fatura No", extracted['fatura_no'], style))
        
        # Kritik notlar
        if critical_notes:
            if style == SummaryStyle.BULLET:
                lines.append("")
                for note in critical_notes[:3]:  # En fazla 3 not
                    lines.append(note)
            else:
                lines.append("")
                lines.append("Uyarılar:")
                for note in critical_notes[:3]:
                    lines.append(f"  • {note}")
        
        # Satır sayısını sınırla
        if len(lines) > settings['max_lines']:
            lines = lines[:settings['max_lines']]
        
        return lines
    
    def _format_line(self, label: str, value: str, style: SummaryStyle) -> str:
        """Satırı stile göre formatla"""
        if style == SummaryStyle.BULLET:
            return f"{label}: **{value}**"
        elif style == SummaryStyle.STRUCTURED:
            return f"{label:12} : {value}"
        else:
            return f"{label}: {value}"
    
    def _format_summary(self, lines: List[str], style: SummaryStyle) -> str:
        """Özet metnini formatla"""
        if style == SummaryStyle.BULLET:
            return "\n".join(lines)
        elif style == SummaryStyle.PARAGRAPH:
            return " ".join(lines)
        else:  # STRUCTURED
            return "\n".join(lines)
    
    def _format_document_type(self, doc_type: str) -> str:
        """Belge türünü formatla"""
        format_map = {
            'fatura': 'Fatura',
            'banka_dekontu': 'Banka Dekontu',
            'maas_bordrosu': 'Maaş Bordrosu',
            'kdv_beyannamesi': 'KDV Beyannamesi',
            'cari_hesap': 'Cari Hesap',
            'genel_tablo': 'Genel Tablo/Rapor',
            'genel_belge': 'Genel Belge',
            'sozlesme': 'Sözleşme',
            'teklif': 'Teklif',
            'dilekce': 'Dilekçe',
            'rapor': 'Rapor',
            'ozgecmis': 'Özgeçmiş'
        }
        return format_map.get(doc_type, doc_type.replace('_', ' ').title())
    
    def _extract_key_points_advanced(self, extracted: Dict, critical_notes: List[str],
                                     keywords: List[str], max_points: int) -> List[str]:
        """Gelişmiş anahtar nokta çıkarımı"""
        points = []
        
        # Önemli bilgilerden noktalar oluştur
        if extracted.get('tutar'):
            points.append(f"💰 {extracted['tutar']}")
        
        if extracted.get('tarih'):
            points.append(f"📅 {extracted['tarih']}")
        
        if extracted.get('firma'):
            points.append(f"🏢 {extracted['firma']}")
        
        if extracted.get('vergi_no'):
            points.append(f"🔢 VN: {extracted['vergi_no']}")
        
        if extracted.get('fatura_no'):
            points.append(f"📋 {extracted['fatura_no']}")
        
        # Kritik notları ekle
        points.extend(critical_notes)
        
        # Anahtar kelimelerden noktalar oluştur (eksikse)
        if len(points) < max_points and keywords:
            keyword_str = ", ".join(keywords[:3])
            points.append(f"🔑 {keyword_str}")
        
        return points[:max_points]
    
    def _calculate_confidence_advanced(self, extracted: Dict, metrics: SummaryMetrics) -> Tuple[int, List[str]]:
        """Gelişmiş güven skoru hesaplama"""
        warnings = []
        
        if not extracted:
            return 20, ["Hiçbir alan çıkarılamadı"]
        
        # Kaç önemli alan bulundu
        fields_found = len(extracted)
        base_confidence = 30 + (fields_found * 10)
        
        # Dil faktörü
        if metrics.language == 'unknown':
            base_confidence -= 10
            warnings.append("Dil tespit edilemedi")
        elif metrics.language != 'tr':
            base_confidence -= 5
            warnings.append(f"Belge dili: {metrics.language.upper()}")
        
        # Metin uzunluğu faktörü
        if metrics.original_length < 50:
            base_confidence -= 20
            warnings.append("Metin çok kısa")
        elif metrics.original_length < 100:
            base_confidence -= 10
            warnings.append("Metin kısa")
        elif metrics.original_length > 1000:
            base_confidence += 10
        
        # Anahtar kelime faktörü
        if metrics.keyword_count > 0:
            base_confidence += min(10, metrics.keyword_count)
        
        confidence = min(100, max(0, base_confidence))
        
        return confidence, warnings
    
    def _generate_suggestions(self, extracted: Dict, metrics: SummaryMetrics) -> List[str]:
        """Öneriler oluştur"""
        suggestions = []
        
        # Eksik alan kontrolü
        expected_fields = ['tarih', 'firma', 'tutar']
        missing = [f for f in expected_fields if f not in extracted]
        
        if missing:
            fields_str = ", ".join(missing)
            suggestions.append(f"Eksik alanlar: {fields_str}")
        
        # Dil önerisi
        if metrics.language != 'tr' and metrics.language != 'en':
            suggestions.append(f"Belge dili: {metrics.language.upper()}, kontrol önerilir")
        
        # Uzunluk önerisi
        if metrics.original_length < 50:
            suggestions.append("Çok kısa metin, özet güvenilirliği düşük")
        
        return suggestions[:3]  # En fazla 3 öneri
    
    def _calculate_readability(self, text: str) -> float:
        """Okunabilirlik skoru hesapla (basitleştirilmiş)"""
        words = text.split()
        sentences = re.split(r'[.!?]+', text)
        
        if not words or not sentences:
            return 0.0
        
        avg_words_per_sentence = len(words) / len(sentences)
        
        # Basit okunabilirlik formülü
        score = max(0, 100 - (avg_words_per_sentence * 2))
        return min(100, score)
    
    def save_summary_to_file(self, summary: str, output_path: str, format_type: str = 'txt') -> bool:
        """Özeti dosyaya kaydet (gelişmiş)"""
        try:
            if format_type.upper() == 'PDF':
                # PDF olarak kaydet
                from reportlab.pdfgen import canvas
                from reportlab.lib.pagesizes import A4
                from reportlab.lib.units import cm
                from reportlab.lib import colors
                
                c = canvas.Canvas(output_path, pagesize=A4)
                width, height = A4
                
                y = height - 2*cm
                line_height = 0.5*cm
                
                # Başlık
                c.setFont("Helvetica-Bold", 16)
                c.setFillColor(colors.HexColor('#2E75B6'))
                c.drawString(2*cm, y, "BELGE ÖZETİ")
                y -= line_height * 2
                
                # İçerik
                c.setFont("Helvetica", 11)
                c.setFillColor(colors.black)
                
                for line in summary.split('\n'):
                    if line.strip():
                        if y < line_height + 1*cm:
                            c.showPage()
                            y = height - 2*cm
                            c.setFont("Helvetica", 11)
                        
                        # Madde işaretleri için girinti
                        if line.startswith('•'):
                            c.drawString(2.5*cm, y, line)
                        else:
                            c.drawString(2*cm, y, line)
                        
                        y -= line_height
                
                # Alt bilgi
                c.setFont("Helvetica", 8)
                c.setFillColor(colors.grey)
                c.drawString(2*cm, 1*cm, f"Oluşturulma: {datetime.datetime.now().strftime('%d.%m.%Y %H:%M')}")
                
                c.save()
                
            elif format_type.upper() == 'WORD':
                # Word olarak kaydet
                from docx import Document
                from docx.shared import Pt, RGBColor
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                
                doc = Document()
                
                # Başlık
                title = doc.add_heading('BELGE ÖZETİ', 0)
                title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = title.runs[0]
                run.font.size = Pt(16)
                run.font.color.rgb = RGBColor(46, 117, 182)
                
                # İçerik
                for line in summary.split('\n'):
                    if line.strip():
                        if line.startswith('•'):
                            doc.add_paragraph(line, style='List Bullet')
                        elif line.startswith('**') and line.endswith('**'):
                            doc.add_heading(line.strip('*'), level=2)
                        else:
                            doc.add_paragraph(line)
                
                doc.save(output_path)
                
            else:
                # TXT olarak kaydet
                with open(output_path, 'w', encoding='utf-8') as f:
                    f.write(summary)
            
            logger.info(f"✅ Özet kaydedildi: {output_path}")
            return True
            
        except Exception as e:
            logger.error(f"❌ Özet kaydetme hatası: {e}")
            return False


# ========== KULLANIM KOLAYLIĞI FONKSİYONLARI ==========

def summarize_document(text: str, document_type: str, 
                      length: str = "orta", style: str = "madde") -> Dict:
    """
    Belgeyi özetle
    Returns: özet bilgileri
    """
    length_map = {
        'kısa': SummaryLength.SHORT,
        'orta': SummaryLength.MEDIUM,
        'detaylı': SummaryLength.DETAILED
    }
    
    style_map = {
        'madde': SummaryStyle.BULLET,
        'paragraf': SummaryStyle.PARAGRAPH,
        'yapısal': SummaryStyle.STRUCTURED
    }
    
    summarizer = DocumentSummarizer(
        length=length_map.get(length, SummaryLength.MEDIUM),
        style=style_map.get(style, SummaryStyle.BULLET)
    )
    
    result = summarizer.summarize(text, document_type)
    
    return {
        'summary': result.summary,
        'key_points': result.key_points,
        'word_count': result.word_count,
        'confidence': result.confidence,
        'document_type': result.document_type,
        'metrics': {
            'original_length': result.metrics.original_length,
            'compression_ratio': f"%{result.metrics.compression_ratio*100:.1f}",
            'language': result.metrics.language,
            'keywords': result.metrics.important_terms
        },
        'warnings': result.warnings,
        'suggestions': result.suggestions
    }

def get_summary_stats(result: Dict) -> str:
    """
    Özet istatistiklerini formatla
    """
    stats = f"""📊 **ÖZET İSTATİSTİKLERİ**

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
📄 **Metin:** {result['word_count']} kelime
🔤 **Dil:** {result['metrics']['language'].upper()}
📦 **Sıkıştırma:** {result['metrics']['compression_ratio']}
⭐ **Güven:** %{result['confidence']}

🔑 **Anahtar Kelimeler:**
"""
    for kw in result['metrics']['keywords'][:5]:
        stats += f"  • {kw}\n"
    
    if result['warnings']:
        stats += "\n⚠️ **Uyarılar:**\n"
        for w in result['warnings']:
            stats += f"  • {w}\n"
    
    return stats


# ========== TEST FONKSİYONU ==========
if __name__ == "__main__":
    print("🔧 Belge Özetleme Modülü Test Ediliyor...")
    print("=" * 60)
    
    test_text = """
    FATURA
    Fatura No: INV2024001
    Tarih: 15.03.2024
    Firma: ABC Ltd. Şti.
    Vergi No: 1234567890
    Tutar: 1.250,50 TL
    KDV: %20
    Toplam: 1.500,60 TL
    IBAN: TR12 0001 0004 6796 3186 2350 01
    
    Mal/Hizmet: Danışmanlık hizmeti
    Miktar: 10 saat
    Birim Fiyat: 125,05 TL
    
    Not: Ödeme vadesi 15.04.2024
    """
    
    # Farklı stillerde test
    summarizer = DocumentSummarizer(length=SummaryLength.DETAILED, style=SummaryStyle.BULLET)
    result = summarizer.summarize(test_text, 'fatura')
    
    print(f"\n📋 **ÖZET (Madde İşaretli):**")
    print(result.summary)
    
    print(f"\n🔑 **Anahtar Noktalar:**")
    for point in result.key_points:
        print(f"  • {point}")
    
    print(f"\n📊 **Güven Skoru:** %{result.confidence}")
    
    if result.warnings:
        print(f"\n⚠️ **Uyarılar:**")
        for warning in result.warnings:
            print(f"  • {warning}")
    
    if result.suggestions:
        print(f"\n💡 **Öneriler:**")
        for suggestion in result.suggestions:
            print(f"  • {suggestion}")
    
    print(f"\n📈 **Metrikler:**")
    print(f"  • Orijinal: {result.metrics.original_length} kelime")
    print(f"  • Özet: {result.metrics.summary_length} kelime")
    print(f"  • Sıkıştırma: %{result.metrics.compression_ratio*100:.1f}")
    print(f"  • Dil: {result.metrics.language.upper()}")
    
    print("=" * 60)
    print("✅ Belge Özetleme Modülü hazır!")